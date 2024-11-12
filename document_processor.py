import os
import uuid
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import logging

class DocumentProcessor:
    def __init__(self, save_dir, chunk_size, chunk_overlap, logging_enabled=False):
        self.save_dir = save_dir
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logging_enabled = logging_enabled
        os.makedirs(self.save_dir, exist_ok=True)

        # Configure logging based on the logging_enabled flag
        self.logger = logging.getLogger(__name__)
        if self.logging_enabled:
            logging.basicConfig(level=logging.DEBUG)
        else:
            self.logger.disabled = True

    def toggle_logging(self, enable_logging):
        """Enable or disable logging dynamically."""
        self.logging_enabled = enable_logging
        self.logger.disabled = not enable_logging

    def extract_text_from_pdf(self, file):
        """Extract text from PDF files."""
        try:
            temp_path = os.path.join(self.save_dir, 'temp.pdf')
            file.save(temp_path)
            
            reader = PdfReader(temp_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            os.remove(temp_path)
            return text
        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {str(e)}")
            return None

    def extract_text_from_docx(self, file):
        """Extract text from DOCX files."""
        try:
            temp_path = os.path.join(self.save_dir, 'temp.docx')
            file.save(temp_path)
            
            doc = Document(temp_path)
            text = ' '.join([para.text for para in doc.paragraphs])
            
            os.remove(temp_path)
            return text
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {str(e)}")
            return None

    def extract_text_from_file(self, uploaded_file):
        """Extract text from various file types."""
        try:
            filename = uploaded_file.filename.lower()
            self.logger.debug(f"Processing file: {filename}")
            
            if filename.endswith('.txt'):
                return uploaded_file.read().decode('utf-8')
            elif filename.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
                return ' '.join(df.astype(str).agg(' '.join, axis=1).tolist())
            elif filename.endswith('.docx'):
                return self.extract_text_from_docx(uploaded_file)
            elif filename.endswith('.pdf'):
                return self.extract_text_from_pdf(uploaded_file)
            
            self.logger.error(f"Unsupported file type: {filename}")
            return None
        except Exception as e:
            self.logger.error(f"Error in extract_text_from_file: {str(e)}")
            return None

    def save_text_to_file(self, text, filename):
        """Save extracted text to a file."""
        try:
            base_name = os.path.splitext(filename)[0]
            unique_filename = f"{uuid.uuid4()}_{base_name}.txt"
            file_path = os.path.join(self.save_dir, unique_filename)
            
            self.logger.debug(f"Saving text to file: {file_path}")
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return file_path
        except Exception as e:
            self.logger.error(f"Error saving text to file: {str(e)}")
            return None

    def load_documents(self, text, filename):
        """Create a LangChain document directly from text."""
        try:
            doc = LangchainDocument(
                page_content=text,
                metadata={"source": filename}
            )
            self.logger.debug(f"Created document with metadata: {doc.metadata}")
            return [doc]
        except Exception as e:
            self.logger.error(f"Error creating document: {str(e)}")
            return None

    def split_documents(self, documents):
        """Split documents into chunks."""
        try:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap
            )
            return text_splitter.split_documents(documents)
        except Exception as e:
            self.logger.error(f"Error splitting documents: {str(e)}")
            return None

    def cleanup_documents(self):
        """Remove files in the SAVE_DIR after processing."""
        try:
            for filename in os.listdir(self.save_dir):
                if filename != 'temp.pdf' and filename != 'temp.docx':
                    file_path = os.path.join(self.save_dir, filename)
                    os.remove(file_path)
        except Exception as e:
            self.logger.error(f"Error cleaning up documents: {str(e)}")

    def process_file(self, uploaded_file):
        """Process a file from upload to splitting."""
        try:
            self.logger.info(f"Starting to process file: {uploaded_file.filename}")
            
            uploaded_file.seek(0)
            
            text = self.extract_text_from_file(uploaded_file)
            if not text:
                self.logger.error("Failed to extract text from file")
                return None

            documents = self.load_documents(text, uploaded_file.filename)
            if not documents:
                self.logger.error("Failed to create document")
                return None

            split_docs = self.split_documents(documents)
            if not split_docs:
                self.logger.error("Failed to split documents")
                return None

            return split_docs
            
        except Exception as e:
            self.logger.error(f"Error in process_file: {str(e)}")
            return None
